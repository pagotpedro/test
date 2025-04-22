
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Leitura do Excel original
df = pd.read_excel("Relatorio de Conversão.xlsx")

# Função para formatar planilha e gerar arquivos auxiliares
def formatar_planilha(df):
    if 'Data_5' in df.columns:
        df.rename(columns={'Data_5': 'Data do TVI'}, inplace=True)

    for col in ['CNPJ ou CPF', 'CNPJ ou CPF_2']:
        if col in df.columns:
            df[col] = df[col].astype(str).str.extract(r'(\d+)')[0].fillna('').str.zfill(11)

    for col in ['Data', 'Data do TVI']:
        if col in df.columns:
            df[col] = pd.to_datetime(df[col], errors='coerce')
            df[col] = df[col].dt.strftime('%d/%m/%Y')

    df['Data do TVI_dt'] = pd.to_datetime(df['Data do TVI'], format='%d/%m/%Y', errors='coerce')
    df['BC + 50%'] = pd.to_numeric(df['Valor do Produto'], errors='coerce').fillna(0) * 1.5

    def calcular_aliquota(data):
        if pd.isnull(data): return None
        if data < datetime(2023, 3, 31): return 0.18
        elif data < datetime(2024, 2, 19): return 0.20
        elif data < datetime(2025, 3, 31): return 0.22
        return 0.23

    df['Aliq Interna'] = df['Data do TVI_dt'].map(calcular_aliquota)
    df['Valor do ICMS'] = pd.to_numeric(df['Valor do ICMS'], errors='coerce').fillna(0)
    df['ICMS Débito'] = df['BC + 50%'] * df['Aliq Interna'] - df['Valor do ICMS']
    df['Aliq Interna'] = df['Aliq Interna'].map(lambda x: f"{x*100:.0f}%" if pd.notnull(x) else "")
    df.drop(columns=['Data do TVI_dt'], inplace=True)

    for col in ['Base de Cálculo do ICMS ST', 'Valor do ICMS ST']:
        if col in df.columns:
            df.drop(columns=[col], inplace=True)

    colunas_financeiras = ['Valor do Produto', 'Base de Cálculo ICMS', 'Valor do ICMS', 'BC + 50%', 'ICMS Débito']
    for col in colunas_financeiras:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)

    df = df[df['Valor Débito TVI'] != 0]

    df_original = df.copy()
    soma = df[colunas_financeiras].sum(numeric_only=True)
    icms_debito_sem_multa = soma.get('ICMS Débito', 0)
    multa = icms_debito_sem_multa / 2
    total_com_multa = icms_debito_sem_multa + multa

    linha_soma = {col: soma[col] if col in soma else '' for col in df.columns}
    linha_multa = {col: multa if col == 'ICMS Débito' else '' for col in df.columns}
    linha_total_multa = {col: total_com_multa if col == 'ICMS Débito' else '' for col in df.columns}
    df = pd.concat([df, pd.DataFrame([linha_soma, linha_multa, linha_total_multa])], ignore_index=True)

    colunas_novas = ['BC + 50%', 'Aliq Interna', 'ICMS Débito']
    if 'Valor da NFe' in df.columns:
        cols = list(df.columns)
        idx = cols.index('Valor da NFe') + 1
        for col in reversed([c for c in colunas_novas if c in cols]):
            cols.insert(idx, cols.pop(cols.index(col)))
        df = df[cols]

    nome_razao = df_original['Razão_4'].dropna().unique()[0]
    resumo = pd.DataFrame({
        'Descrição': [
            f'Quadro Resumo - {nome_razao}', '',
            'Valor total dos produtos',
            'BC Aplicada - Base de Cálculo + 50%',
            'ICMS Débito = Alíquota x BC',
            'Crédito de ICMS destacado em NF-e',
            'Valor ICMS a recolher',
            'Multa de 50%',
            'Total Devido (ICMS a recolher + Multa de 50%)'
        ],
        'Valor': [
            '', '',
            soma.get('Valor do Produto', 0),
            soma.get('BC + 50%', 0),
            icms_debito_sem_multa + soma.get('Valor do ICMS', 0),
            soma.get('Valor do ICMS', 0),
            icms_debito_sem_multa,
            multa,
            total_com_multa
        ]
    })

    gfis_df = pd.DataFrame()
    gfis_df['Inscrição Renavam_3'] = df_original['Inscrição Renavam_3'].astype(str)
    gfis_df['CNPJ ou CPF_2'] = df_original['CNPJ ou CPF_2'].astype(str)
    gfis_df['Data do TVI'] = df_original['Data do TVI'].astype(str)
    gfis_df['ICMS Débito'] = df_original['ICMS Débito'].astype(float).round(2)

    return df, resumo, gfis_df, nome_razao

df_formatado, resumo_df, gfis_df, razao = formatar_planilha(df)

# Nomes de arquivos
nome_limpo = razao.replace("/", "-").upper()
planilha_calc = f"Planilha de Cálculo - {nome_limpo}.xlsx"
planilha_gfis = f"GFIS - {nome_limpo}.xlsx"
word_saida = f"Relatório de Conversão - {nome_limpo}.docx"

# Exportar planilha formatada e quadro resumo
with pd.ExcelWriter(planilha_calc, engine='xlsxwriter') as writer:
    df_formatado.to_excel(writer, index=False, sheet_name="Planilha Formatada")
    resumo_df.to_excel(writer, index=False, sheet_name="Quadro Resumo", startrow=1)

# Exportar GFIS
gfis_df.to_excel(planilha_gfis, index=False)

# Gerar Word
doc = Document("TEMPLATE_NORMAL_ANTECIPADO.docx")
for p in doc.paragraphs:
    if "#INSCRICAO" in p.text:
        p.text = p.text.replace("#INSCRICAO", str(df['Inscrição Renavam_3'].dropna().iloc[0]))
    if "#CNPJ" in p.text:
        p.text = p.text.replace("#CNPJ", str(df['CNPJ ou CPF_2'].dropna().iloc[0]))
    if "#RAZAO" in p.text:
        p.text = p.text.replace("#RAZAO", razao)
    if "#DATA" in p.text:
        p.text = p.text.replace("#DATA", datetime.today().strftime("%d/%m/%Y"))

for p in doc.paragraphs:
    if "#QUADRORESUMO" in p.text:
        p.text = ""
        p_element = p._element
        parent = p_element.getparent()
        idx = parent.index(p_element)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Descrição"
        hdr_cells[1].text = "Valor"
        for cell in hdr_cells:
            for par in cell.paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = par.runs[0]
                run.bold = True
                run.font.size = Pt(10)
        resumo_limpo = resumo_df.iloc[2:].reset_index(drop=True)
        for _, row in resumo_limpo.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row["Descrição"])
            row_cells[1].text = "R$ {:,.2f}".format(row["Valor"]).replace(",", "#").replace(".", ",").replace("#", ".")
        parent.insert(idx, table._element)
        break

doc.save(word_saida)
print("✅ Arquivos gerados com sucesso!")
