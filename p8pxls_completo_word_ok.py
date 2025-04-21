def gerar_docx_quadroresumo_formatado(inscricao, cnpj, razao, resumo_df):
    template_bytes = base64.b64decode(template_base64_valido)
    doc = Document(BytesIO(template_bytes))

    for i, par in enumerate(doc.paragraphs):
        if "#INSCRICAO" in par.text:
            par.text = par.text.replace("#INSCRICAO", inscricao)
        if "#CNPJ" in par.text:
            par.text = par.text.replace("#CNPJ", cnpj)
        if "#RAZAO" in par.text:
            par.text = par.text.replace("#RAZAO", razao)
        if "#DATA" in par.text:
            par.text = par.text.replace("#DATA", datetime.now().strftime('%d/%m/%Y'))
        if "#QUADRORESUMO" in par.text:
            p = par._element
            parent = p.getparent()
            idx = parent.index(p)
            parent.remove(p)

            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"

            # Cabeçalho
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Descrição"
            hdr_cells[1].text = "Valor"
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(10)

            # Valores formatados
            for _, row in resumo_df.iterrows():
                if pd.notna(row["Descrição"]) and row["Descrição"] != "":
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row["Descrição"])
                    valor_formatado = "R$ {:,.2f}".format(row["Valor"]).replace(",", "#").replace(".", ",").replace("#", ".")
                    row_cells[1].text = valor_formatado

            parent.insert(idx, table._element)

    output = BytesIO()
    doc.save(output)
    return output



import streamlit as st
import pandas as pd
from io import BytesIO
import base64
from docx import Document

# Template Word embutido
