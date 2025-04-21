def gerar_docx_com_quadro_resumo(inscricao, cnpj, razao, resumo_df):
    template_bytes = base64.b64decode(template_base64_valido)
    doc = Document(BytesIO(template_bytes))

    # Substituição de campos simples
    for i, par in enumerate(doc.paragraphs):
        if "#INSCRICAO" in par.text:
            par.text = par.text.replace("#INSCRICAO", inscricao)
        if "#CNPJ" in par.text:
            par.text = par.text.replace("#CNPJ", cnpj)
        if "#RAZAO" in par.text:
            par.text = par.text.replace("#RAZAO", razao)
        if "#DATA" in par.text:
            par.text = par.text.replace("#DATA", datetime.now().strftime('%d/%m/%Y'))

        # Inserir tabela do quadro resumo no local do marcador
        if "#QUADRORESUMO" in par.text:
            p = par._element
            parent = p.getparent()
            idx = parent.index(p)
            parent.remove(p)

            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Descrição"
            hdr_cells[1].text = "Valor"

            for _, row in resumo_df.iterrows():
                if pd.notna(row["Descrição"]) and row["Descrição"] != "":
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row["Descrição"])
                    row_cells[1].text = f'R$ {row["Valor"]:,.2f}'.replace('.', '#').replace(',', '.').replace('#', ',')

            # Adiciona a nova tabela no local da antiga parágrafo
            parent.insert(idx, table._element)

    output = BytesIO()
    doc.save(output)
    return output



import streamlit as st
import pandas as pd
from io import BytesIO
import base64
from docx import Document

# Template Word embutido
